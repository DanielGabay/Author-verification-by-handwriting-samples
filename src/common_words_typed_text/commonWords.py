import docx
from docx.enum.section import WD_ORIENT
import string
import sys
import os
from collections import Counter 
import numpy as np
import matplotlib.pyplot as plt

# constatns
DATA_PATH = "data/"
RESULTS_PATH = "results.docx"
K_TOP_WORDS = 15

def get_text_as_list(filename):
    '''
    returns file text as a list, after removing all punctuations
    '''
    print("reading file:", filename)
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    exclude = set(string.punctuation)
    s = '\n'.join(fullText)
    s = ''.join(ch for ch in s if ch not in exclude)
    return s.split()

def get_k_common_words(filepath):
    text_list = get_text_as_list(filepath)
    # Pass the split_it list to instance of Counter class. 
    counter = Counter(text_list) 
    
    # most_common() produces k frequently encountered 
    # input values and their respective counts. 
    return counter.most_common(K_TOP_WORDS)

def fill_row_data(row_cells, top_words):
    '''
    fills top_words in a row cells
    '''
    print("filling row data")
    cell_num = 1
    for word in top_words:
        res = str(word[0]) + "\n" + str(word[1])
        row_cells[cell_num].text = res
        cell_num += 1

def set_doc_landscape(document):
    '''
    set document layout as landscape
    '''
    sections = document.sections
    for section in sections:
        # change orientation to landscape
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

def create_table():
    # create a table inside a result document
    print("creating ", RESULTS_PATH , " document")
    document = docx.Document()
    table = document.add_table(rows=0, cols=K_TOP_WORDS+1, style='Table Grid')

    # go through each file in DATA_PATH
    for root, dirs, files in os.walk(DATA_PATH):
        for file in files:
            print("-------------------------")
            filepath = DATA_PATH + file
            # extract file top words
            top_words = get_k_common_words(filepath)
            # fill top words data in a new row
            row_cells = table.add_row().cells
            row_cells[0].text = str(file)
            fill_row_data(row_cells, top_words)
  
    print("-------------------------")
    set_doc_landscape(document)
    document.save(RESULTS_PATH)
    print(RESULTS_PATH , "saved successfully!")

def get_words_freq_list(num_of_words=K_TOP_WORDS):
    all_words = []
    # go through each file in DATA_PATH
    for root, dirs, files in os.walk(DATA_PATH):
        for file in files:
            print("-------------------------")
            filepath = DATA_PATH + file
            text_list = get_text_as_list(filepath)
            # all words is a matrix that each row contains a set (unique) words
            # from each file
            all_words.append(set(text_list))
    
    words_freq_list = []
    # for each word in a file, check how many files contains that words
    for row in all_words:
        for word in row:
            # to avoid duplicates, check if the word was already added
            # word[::-1] -> reverse the word for currect output
            if not any(word[::-1] in sublist for sublist in words_freq_list) :
                # append word and its freq
                words_freq_list.append([count_in_files(row,all_words,word), word[::-1]])

    # sort in descending order
    words_freq_list = sorted(words_freq_list, key = lambda x: x[0], reverse=True)
    # remain only the most freq words
    words_freq_list = words_freq_list[:num_of_words]
    return words_freq_list

def rever(strings):
    return [x[::-1] for x in strings]

def check_most_common_words_intersection(print_intesection=True ):
    common_words_to_compare = ['של', 'לא' ,'את', 'על', 'לסיכום', 'כי', 'זה', 'זו', 'גם', 'לדעתי']
    all_words = []
    # go through each file in DATA_PATH
    for root, dirs, files in os.walk(DATA_PATH):
        for file in files:
            filepath = DATA_PATH + file
            words_freq_list = get_k_common_words(filepath)
            # sort in descending order
            words_freq_list = sorted(words_freq_list, key = lambda x: x[0], reverse=True)
            # remain only the most freq words
            words_freq_list = words_freq_list[:K_TOP_WORDS]
            words, freq = zip(*words_freq_list)
            words = [x for x in words if x in common_words_to_compare]

            # all words is a matrix that each row contains a set (unique) words
            # from each file
            all_words.append(rever(set(words)))
    
    all_intersect = True
    for i in range(len(all_words)):
        for j in range(i+1, len(all_words)):
            intersection = set(all_words[i]).intersection(set(all_words[j]))
            if(len(intersection) == 0):
                print("files: " + str(i+1) + ".docx " + str(j+1) + ".docx not intersect!")
                all_intersect = False
            if(print_intesection):
                print("Check files: " + str(i+1) + ".docx " + str(j+1) +".docx")
                print(intersection)
                print("--------------------")
    if(all_intersect):
        print("All files have intersect of common words!")

def create_graph():
    words_freq_list = get_words_freq_list()
    # split into words array and freq array
    freq, words = zip(*words_freq_list)

    indices = np.arange(len(words_freq_list))
    plt.bar(indices, freq, color='b')
    plt.xticks(indices, words, rotation='vertical')
    plt.tight_layout()
    plt.show()


def count_in_files(row,all_words, word):
    count = 1
    for r in all_words:
        if r != row and word in r:
            count +=1
    return count
            
def main(): 
    if(len(sys.argv) < 2):
        print("Usage: python words.py [-t] [-g] [-i <t/f>]")
        return
    flag = sys.argv[1]
    if(flag == "-t"):
        create_table()

    elif(flag == "-g"):
        create_graph()

    elif(flag == "-i"):
        if len(sys.argv) < 3:
            print("Usage: python words.py [-t] [-g] [-i <t/f>]")
            return
        if sys.argv[2] == 't':
            check_most_common_words_intersection(True)
        elif sys.argv != 'f':
            check_most_common_words_intersection(False)
        else:
            print("Usage: python words.py [-t] [-g] [-i <t/f>]")
            return



if __name__ == "__main__":
    main()
